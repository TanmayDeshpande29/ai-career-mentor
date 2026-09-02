from io import BytesIO
from pathlib import Path

from langchain_core.documents import Document
from pypdf import PdfReader
from docx import Document as DocxDocument


class ResumeDocumentLoader:

    @staticmethod
    def load_from_bytes(
        file_content: bytes,
        file_name: str,
        content_type: str | None = None,
    ) -> list[Document]:

        extension = Path(file_name).suffix.lower()

        if extension == ".pdf":
            return ResumeDocumentLoader._load_pdf_bytes(
                file_content,
                file_name,
            )

        if extension == ".docx":
            return ResumeDocumentLoader._load_docx_bytes(
                file_content,
                file_name,
            )

        raise ValueError(
            f"Unsupported resume format: {extension}"
        )

    @staticmethod
    def _load_pdf_bytes(
        file_content: bytes,
        file_name: str,
    ) -> list[Document]:

        reader = PdfReader(
            BytesIO(file_content)
        )

        documents = []

        for page_number, page in enumerate(
            reader.pages,
            start=1,
        ):

            text = page.extract_text() or ""
            text = text.strip()

            if not text:
                continue

            documents.append(
                Document(
                    page_content=text,
                    metadata={
                        "source": file_name,
                        "file_type": "pdf",
                        "page": page_number,
                    },
                )
            )

        return documents

    @staticmethod
    def _load_docx_bytes(
        file_content: bytes,
        file_name: str,
    ) -> list[Document]:

        document = DocxDocument(
            BytesIO(file_content)
        )

        paragraphs = [
            paragraph.text.strip()
            for paragraph in document.paragraphs
            if paragraph.text.strip()
        ]

        text = "\n".join(paragraphs)

        if not text:
            return []

        return [
            Document(
                page_content=text,
                metadata={
                    "source": file_name,
                    "file_type": "docx",
                    "page": 1,
                },
            )
        ]